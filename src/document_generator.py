import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(markdown_text: str, output_path: str):
    """
    Converts simple Markdown to a stylized DOCX document.
    Handles headers, bullet points, and basic bold text.
    """
    doc = Document()
    
    # Optional: Set a nice default font for the whole document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            # Clean bold markers inside headers
            text = text.replace('**', '')
            
            heading = doc.add_heading(text, level=min(level, 4))
            
            # Center the top-level header (usually Candidate Name)
            if level == 1:
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            continue

        # Bullet Points
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
            continue
            
        # Normal Paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, line)

    doc.save(output_path)

def _add_formatted_text(paragraph, text: str):
    """
    Parses **bold** text within a paragraph and adds it to the python-docx paragraph.
    """
    # Split by ** to find bold sections
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # It's bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Normal text
            paragraph.add_run(part)
