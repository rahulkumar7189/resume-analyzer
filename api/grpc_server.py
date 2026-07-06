import os
os.environ["CUDA_VISIBLE_DEVICES"] = ""

import sys
import pathlib
import time
import hashlib
import grpc
import re
import shutil
import json
from concurrent import futures

parent_dir = str(pathlib.Path(__file__).resolve().parent.parent)
if parent_dir not in sys.path:
    sys.path.append(parent_dir)

def load_env_vars():
    env_path = pathlib.Path(__file__).resolve().parent.parent / "ats-web" / ".env.local"
    if env_path.exists():
        print(f"[grpc_server] Loading env vars from {env_path}")
        with open(env_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    key, value = line.split("=", 1)
                    os.environ.setdefault(key.strip(), value.strip())

load_env_vars()

# Ensure temp and upload directories exist
temp_dir = pathlib.Path(parent_dir) / "temp"
temp_dir.mkdir(exist_ok=True)
upload_dir = pathlib.Path(parent_dir) / "ats-web" / "public" / "uploads"
upload_dir.mkdir(parents=True, exist_ok=True)

# Generate stubs if they don't exist yet (we'll do this in a startup script, but assuming they exist here)
try:
    sys.path.append(str(pathlib.Path(parent_dir) / "protos"))
    import ats_pb2
    import ats_pb2_grpc
except ImportError as e:
    print(f"[grpc_server] Protobuf stubs not found. Run grpcio-tools to generate them. Error: {e}")
    sys.exit(1)

from api.redis_client import get_cached_result
from api.celery_app import celery_app
from api.tasks import process_resume_task, autofix_resume_task, suggest_edits_task, compile_pdf_task

_PIPELINE_AVAILABLE = False
try:
    from src.extractor import extract_text
    from src.llm_engine import suggest_skill_integration, client as groq_client
    _PIPELINE_AVAILABLE = True
    print("[grpc_server] ML pipeline imported successfully.")
except Exception as e:
    print(f"[grpc_server] WARNING: Pipeline import failed: {e}")

# ── File parsing helpers ──────────────────────────────────────────────────────

def _parse_pdf(path: str) -> str:
    try:
        return extract_text(path)
    except Exception:
        try:
            import fitz
            doc = fitz.open(path)
            return "\n".join(page.get_text("text") for page in doc)
        except Exception as e:
            raise Exception(f"PDF extraction failed: {e}")

def _parse_docx(path: str) -> str:
    try:
        import docx
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        raise Exception("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {e}")

def parse_uploaded_file(file_obj, temp_path: str) -> str:
    filename = (file_obj.filename or "").lower()
    if filename.endswith(".txt"):
        with open(temp_path, "rb") as f:
            content = f.read()
        return content.decode("utf-8", errors="ignore")

    if filename.endswith(".pdf"):
        return _parse_pdf(temp_path)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        return _parse_docx(temp_path)
    else:
        try:
            with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Unsupported file: {filename}")

_EMAIL_RE = re.compile(r'[\w.+-]+@[\w.-]+\.\w{2,}')
def _extract_email(text: str) -> str:
    m = _EMAIL_RE.search(text)
    return m.group(0) if m else "extracted@email.com"

def extract_candidate_metadata(text: str, filename: str) -> dict:
    email = _extract_email(text)
    fallback_name = (
        filename.replace('.pdf', '').replace('.docx', '').replace('.txt', '')
                .replace('_', ' ').replace('-', ' ').title()
    )
    result = {
        "name": fallback_name,
        "email": email,
        "social_links": {
            "github": "", "linkedin": "", "leetcode": "",
            "codeforces": "", "hackerrank": "",
        },
    }

    if not _PIPELINE_AVAILABLE:
        return result

    try:
        snippet = text[:1500]
        prompt = (
            "You are an ATS metadata parser. Read this resume snippet and extract the candidate's "
            "full name and any online profile URLs (GitHub, LinkedIn, LeetCode, CodeForces, HackerRank).\n"
            "Respond ONLY with raw JSON (no markdown):\n"
            '{"name": "Name", "github": "https://github.com/...", "linkedin": "https://linkedin.com/in/...", '
            '"leetcode": "https://leetcode.com/...", "codeforces": "https://codeforces.com/profile/...", '
            '"hackerrank": "https://hackerrank.com/..."}\n'
            "Leave missing URLs as empty strings. Only return complete URLs.\n\n"
            f"Resume snippet:\n{snippet}"
        )
        user_hash = hashlib.sha256(email.encode("utf-8", errors="ignore")).hexdigest()

        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.1,
            max_tokens=150,
            response_format={"type": "json_object"},
            messages=[{"role": "user", "content": prompt}],
            user=user_hash,
        )
        msg = completion.choices[0].message
        if hasattr(msg, "refusal") and msg.refusal:
            return result
        data = json.loads(msg.content)

        extracted_name = data.get("name", "").strip()
        if extracted_name and len(extracted_name) < 50 and "here is" not in extracted_name.lower():
            result["name"] = extracted_name

        for key in ("github", "linkedin", "leetcode", "codeforces", "hackerrank"):
            result["social_links"][key] = data.get(key, "").strip()
    except Exception as e:
        print(f"[grpc_server] Metadata LLM call failed: {e}")

    return result

class AtsServiceServicer(ats_pb2_grpc.AtsServiceServicer):
    def AnalyzeResume(self, request, context):
        if not _PIPELINE_AVAILABLE:
            context.abort(grpc.StatusCode.UNAVAILABLE, "Pipeline not available")

        filename = request.file_name or "resume.pdf"
        temp_path = str(temp_dir / f"tmp_{int(time.time())}_{filename}")

        try:
            with open(temp_path, "wb") as f:
                f.write(request.file_chunk)

            # Re-use the existing file parsing logic
            class MockUploadFile:
                def __init__(self, filename):
                    self.filename = filename

            mock_upload = MockUploadFile(filename)
            resume_text = parse_uploaded_file(mock_upload, temp_path)
            
            if not resume_text.strip():
                context.abort(grpc.StatusCode.INVALID_ARGUMENT, "Uploaded resume contains no readable text.")

            metadata = extract_candidate_metadata(resume_text, filename)
            if request.candidate_email and request.candidate_email.strip():
                metadata['email'] = request.candidate_email.strip()

            cache_key = hashlib.sha256(f"{resume_text}_{request.job_description}".encode("utf-8")).hexdigest()
            cached_result = get_cached_result(cache_key)
            if cached_result:
                import json
                return ats_pb2.AnalyzeResponse(
                    status="success",
                    task_id="",
                    detail=json.dumps(cached_result)
                )

            # Save PDF locally
            unique_name = f"{int(time.time())}_{filename}"
            save_path = upload_dir / unique_name
            with open(save_path, "wb") as f:
                f.write(request.file_chunk)
            resume_url = f"/uploads/{unique_name}"

            task = process_resume_task.delay(
                resume_text=resume_text,
                resume_filename=filename,
                job_description=request.job_description,
                metadata=metadata,
                resume_url=resume_url,
                job_id=request.job_id,
                recruiter_id=request.recruiter_id,
                cache_key=cache_key
            )
            return ats_pb2.AnalyzeResponse(status="queued", task_id=task.id)

        except Exception as e:
            print(f"[grpc_server] Critical error: {e}")
            import traceback; traceback.print_exc()
            context.abort(grpc.StatusCode.INTERNAL, str(e))
        finally:
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass

    def GetTaskStatus(self, request, context):
        task = celery_app.AsyncResult(request.task_id)
        if task.state in ['PENDING', 'STARTED']:
            return ats_pb2.TaskStatusResponse(status="processing", state=task.state)
        elif task.state == 'SUCCESS':
            result = task.result
            if result["status"] == "success":
                resp = ats_pb2.TaskStatusResponse(status="success", state="SUCCESS")
                if "new_resume_url" in result:
                    resp.new_resume_url = result["new_resume_url"]
                elif "edits" in result:
                    import json
                    resp.data_inserted_json = json.dumps(result["edits"])
                else:
                    import json
                    resp.data_inserted_json = json.dumps([result["data"]])
                return resp
            else:
                return ats_pb2.TaskStatusResponse(status="error", state="SUCCESS", detail=result.get("detail", "Unknown error"))
        elif task.state == 'FAILURE':
            return ats_pb2.TaskStatusResponse(status="error", state="FAILURE", detail=str(task.info))
        else:
            return ats_pb2.TaskStatusResponse(status="processing", state=task.state)

    def SuggestSkill(self, request, context):
        if not _PIPELINE_AVAILABLE:
            context.abort(grpc.StatusCode.UNAVAILABLE, "Pipeline not available")
        suggestion = suggest_skill_integration(request.resume_text, request.skill)
        return ats_pb2.SuggestSkillResponse(status="success", suggestion=suggestion)

    def AutofixResume(self, request, context):
        improvement_tips = [{"original": tip.original, "suggested": tip.suggested} for tip in request.improvement_tips]
        task = autofix_resume_task.delay(
            resume_url=request.resume_url,
            job_description=request.job_description,
            missing_keywords=list(request.missing_keywords),
            improvement_tips=improvement_tips,
            output_format=request.output_format
        )
        return ats_pb2.AutofixResponse(status="queued", task_id=task.id)

    def SuggestEdits(self, request, context):
        improvement_tips = [{"original": tip.original, "suggested": tip.suggested} for tip in request.improvement_tips]
        task = suggest_edits_task.delay(
            resume_url=request.resume_url,
            job_description=request.job_description,
            missing_keywords=list(request.missing_keywords),
            improvement_tips=improvement_tips
        )
        return ats_pb2.SuggestEditsResponse(status="queued", task_id=task.id)

    def CompilePdf(self, request, context):
        accepted_edits = [{"original": edit.original, "suggested": edit.suggested} for edit in request.accepted_edits]
        task = compile_pdf_task.delay(
            resume_url=request.resume_url,
            accepted_edits=accepted_edits,
            template=request.template
        )
        return ats_pb2.CompilePdfResponse(status="queued", task_id=task.id)

def serve():
    server = grpc.server(futures.ThreadPoolExecutor(max_workers=10))
    ats_pb2_grpc.add_AtsServiceServicer_to_server(AtsServiceServicer(), server)
    server.add_insecure_port('[::]:50051')
    server.start()
    print("[grpc_server] gRPC Server running on port 50051...")
    server.wait_for_termination()

if __name__ == '__main__':
    serve()
