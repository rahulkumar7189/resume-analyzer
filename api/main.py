# api/main.py – FastAPI entry point for resume parsing and analysis
import os
os.environ["CUDA_VISIBLE_DEVICES"] = ""  # Force CPU to prevent GPU VRAM hogging
import sys
import pathlib
import shutil
import json
import re
import time
import asyncio
import hashlib
from typing import List, Optional

# Add parent directory to sys.path so we can import from `src`
parent_dir = str(pathlib.Path(__file__).resolve().parent.parent)
if parent_dir not in sys.path:
    sys.path.append(parent_dir)

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from fastapi.staticfiles import StaticFiles

# ── Load environment variables from .env.local ────────────────────────────────
def load_env_vars():
    env_path = pathlib.Path(__file__).resolve().parent.parent / "ats-web" / ".env.local"
    if env_path.exists():
        print(f"[api] Loading env vars from {env_path}")
        with open(env_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    key, value = line.split("=", 1)
                    os.environ.setdefault(key.strip(), value.strip())

load_env_vars()

from api.redis_client import get_cached_result, set_cached_result
from api.celery_app import celery_app
from api.tasks import process_resume_task

# ── ML pipeline imports ───────────────────────────────────────────────────────
_PIPELINE_AVAILABLE = False
try:
    from src.extractor import extract_text, extract_skills_section_aware
    from src.scorer import calculate_ats_match
    from src.llm_engine import generate_strict_feedback, ResumeFeedback, client as groq_client
    from src.parser import parse_resume
    from src.quality_checker import check_formatting
    _PIPELINE_AVAILABLE = True
    print("[api] ML pipeline imported successfully.")
except Exception as e:
    print(f"[api] WARNING: Pipeline import failed: {e}")

# ── FastAPI app ───────────────────────────────────────────────────────────────
app = FastAPI(
    title="Intelligent ATS Backend",
    description="FastAPI service with Celery Background Jobs",
    version="3.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount the uploads directory natively via FastAPI to bypass Next.js caching bugs
upload_dir = pathlib.Path(parent_dir) / "ats-web" / "public" / "uploads"
upload_dir.mkdir(parents=True, exist_ok=True)
app.mount("/uploads", StaticFiles(directory=str(upload_dir)), name="uploads")


# ── File parsing helpers ──────────────────────────────────────────────────────

def _parse_pdf(path: str) -> str:
    try:
        return extract_text(path)
    except Exception:
        try:
            import fitz
            doc = fitz.open(path)
            return "\n".join(page.get_text("text") for page in doc)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF extraction failed: {e}")

def _parse_docx(path: str) -> str:
    try:
        import docx
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed. Run: pip install python-docx",
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"DOCX extraction failed: {e}")

def parse_uploaded_file(file: UploadFile, temp_path: str) -> str:
    filename = (file.filename or "").lower()
    if filename.endswith(".txt"):
        content = file.file.read()
        return content.decode("utf-8", errors="ignore")

    with open(temp_path, "wb") as buf:
        shutil.copyfileobj(file.file, buf)

    if filename.endswith(".pdf"):
        return _parse_pdf(temp_path)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        return _parse_docx(temp_path)
    else:
        try:
            with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Unsupported file: {filename}")

# ── Metadata extraction helpers ───────────────────────────────────────────────

_EMAIL_RE = re.compile(r'[\w.+-]+@[\w.-]+\.\w{2,}')

def _extract_email(text: str) -> str:
    m = _EMAIL_RE.search(text)
    return m.group(0) if m else "extracted@email.com"

def extract_candidate_metadata(text: str, filename: str) -> dict:
    email = _extract_email(text)
    fallback_name = (
        filename.replace('.pdf', '').replace('.docx', '').replace('.txt', '')
                .replace('_', ' ').replace('-', ' ').title()
    )
    result = {
        "name": fallback_name,
        "email": email,
        "social_links": {
            "github": "", "linkedin": "", "leetcode": "",
            "codeforces": "", "hackerrank": "",
        },
    }

    if not _PIPELINE_AVAILABLE or not groq_client:
        return result

    try:
        snippet = text[:1500]
        prompt = (
            "You are an ATS metadata parser. Read this resume snippet and extract the candidate's "
            "full name and any online profile URLs (GitHub, LinkedIn, LeetCode, CodeForces, HackerRank).\n"
            "Respond ONLY with raw JSON (no markdown):\n"
            '{"name": "Name", "github": "https://github.com/...", "linkedin": "https://linkedin.com/in/...", '
            '"leetcode": "https://leetcode.com/...", "codeforces": "https://codeforces.com/profile/...", '
            '"hackerrank": "https://hackerrank.com/..."}\n'
            "Leave missing URLs as empty strings. Only return complete URLs.\n\n"
            f"Resume snippet:\n{snippet}"
        )
        user_hash = hashlib.sha256(email.encode("utf-8", errors="ignore")).hexdigest()

        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.1,
            max_tokens=150,
            response_format={"type": "json_object"},
            messages=[{"role": "user", "content": prompt}],
            user=user_hash,
        )
        msg = completion.choices[0].message
        if hasattr(msg, "refusal") and msg.refusal:
            return result
        data = json.loads(msg.content)

        extracted_name = data.get("name", "").strip()
        if extracted_name and len(extracted_name) < 50 and "here is" not in extracted_name.lower():
            result["name"] = extracted_name

        for key in ("github", "linkedin", "leetcode", "codeforces", "hackerrank"):
            result["social_links"][key] = data.get(key, "").strip()
    except Exception as e:
        print(f"[api] Metadata LLM call failed: {e}")

    return result


# ── Routes ────────────────────────────────────────────────────────────────────

class SuggestSkillRequest(BaseModel):
    resume_text: str
    skill: str

@app.post("/api/suggest-skill")
def suggest_skill(req: SuggestSkillRequest):
    if not _PIPELINE_AVAILABLE:
        raise HTTPException(status_code=503, detail="Pipeline not available")
    
    from src.llm_engine import suggest_skill_integration
    suggestion = suggest_skill_integration(req.resume_text, req.skill)
    return {"status": "success", "suggestion": suggestion}

@app.get("/")
def read_root():
    return {
        "status": "online",
        "version": "3.0.0",
        "pipeline_active": _PIPELINE_AVAILABLE,
    }

@app.get("/api/health")
def health_check():
    return {"status": "ok"}


@app.post("/api/analyze")
async def analyze_resume(
    resume: UploadFile = File(...),
    job_description: str = Form(""),
    job_id: Optional[str] = Form(None),
    recruiter_id: Optional[str] = Form(None),
    candidate_email: Optional[str] = Form(None),
    authorization: Optional[str] = Header(None),
):
    """
    Submits a resume for processing to the Celery worker queue.
    """
    print(f"[api] Analyze request queued: {resume.filename}")

    temp_dir = pathlib.Path(parent_dir) / "temp"
    temp_dir.mkdir(exist_ok=True)
    safe_filename = pathlib.Path(resume.filename).name if resume.filename else "upload.pdf"
    temp_path = str(temp_dir / f"tmp_{int(time.time())}_{safe_filename}")

    try:
        resume_text = parse_uploaded_file(resume, temp_path)
        if not resume_text.strip():
            raise HTTPException(status_code=400, detail="Uploaded resume contains no readable text.")

        metadata = extract_candidate_metadata(resume_text, resume.filename)
        
        # Override parsed email with the logged-in user's email if provided
        if candidate_email and candidate_email.strip():
            metadata['email'] = candidate_email.strip()
            
        print(f"[api] Candidate metadata parsed: {metadata['name']} ({metadata['email']})")

        # Cache check using Redis
        cache_key = hashlib.sha256(f"{resume_text}_{job_description}".encode("utf-8")).hexdigest()
        cached_result = get_cached_result(cache_key)
        if cached_result:
            print(f"[api] Redis Cache HIT for {metadata['name']}")
            return {
                "status": "success",
                "db_saved": True,
                "data_inserted": [cached_result]
            }

        # Save PDF locally
        resume_url = ""
        try:
            upload_dir = pathlib.Path(parent_dir) / "ats-web" / "public" / "uploads"
            upload_dir.mkdir(parents=True, exist_ok=True)
            unique_name = f"{int(time.time())}_{safe_filename}"
            save_path = upload_dir / unique_name
            
            resume.file.seek(0)
            with open(save_path, "wb") as buffer:
                shutil.copyfileobj(resume.file, buffer)
            resume_url = f"/uploads/{unique_name}"
        except Exception as e:
            print(f"[api] Storage save skipped: {e}")

        # Dispatch background task to Celery
        task = process_resume_task.delay(
            resume_text=resume_text,
            resume_filename=resume.filename,
            job_description=job_description,
            metadata=metadata,
            resume_url=resume_url,
            job_id=job_id,
            recruiter_id=recruiter_id,
            cache_key=cache_key
        )
        
        return {"status": "queued", "task_id": task.id}

    except HTTPException:
        raise
    except Exception as e:
        print(f"[api] Critical error: {e}")
        import traceback; traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


@app.get("/api/task/{task_id}")
async def get_task_status(task_id: str):
    """
    Check the status of a Celery background task.
    """
    task = celery_app.AsyncResult(task_id)
    if task.state == 'PENDING' or task.state == 'STARTED':
        return {"status": "processing"}
    elif task.state == 'SUCCESS':
        result = task.result
        if result["status"] == "success":
            if "new_resume_url" in result:
                return {"status": "success", "new_resume_url": result["new_resume_url"]}
            if "edits" in result:
                return {"status": "success", "edits": result["edits"]}
            return {"status": "success", "data_inserted": [result["data"]]}
        else:
            return {"status": "error", "detail": result.get("detail", "Unknown error")}
    elif task.state == 'FAILURE':
        return {"status": "error", "detail": str(task.info)}
    else:
        return {"status": "processing", "state": task.state}

from pydantic import BaseModel

class AutofixRequest(BaseModel):
    resume_url: str
    job_description: str
    missing_keywords: list[str]
    improvement_tips: list[dict]
    output_format: str = "tex"

@app.post("/api/autofix")
async def autofix_resume(req: AutofixRequest):
    from api.tasks import autofix_resume_task
    task = autofix_resume_task.delay(
        resume_url=req.resume_url,
        job_description=req.job_description,
        missing_keywords=req.missing_keywords,
        improvement_tips=req.improvement_tips,
        output_format=req.output_format
    )
    return {"status": "queued", "task_id": task.id}

class SuggestEditsRequest(BaseModel):
    resume_url: str
    job_description: str
    missing_keywords: list[str]
    improvement_tips: list[dict]

@app.post("/api/suggest-edits")
async def suggest_edits(req: SuggestEditsRequest):
    if not _PIPELINE_AVAILABLE:
        raise HTTPException(status_code=503, detail="Pipeline not available")
    from api.tasks import suggest_edits_task
    task = suggest_edits_task.delay(
        resume_url=req.resume_url,
        job_description=req.job_description,
        missing_keywords=req.missing_keywords,
        improvement_tips=req.improvement_tips
    )
    return {"status": "queued", "task_id": task.id}

class CompilePdfRequest(BaseModel):
    resume_url: str
    accepted_edits: list[dict]
    template: str = "modern"

@app.post("/api/compile-pdf")
async def compile_pdf(req: CompilePdfRequest):
    if not _PIPELINE_AVAILABLE:
        raise HTTPException(status_code=503, detail="Pipeline not available")
    from api.tasks import compile_pdf_task
    task = compile_pdf_task.delay(
        resume_url=req.resume_url,
        accepted_edits=req.accepted_edits,
        template=req.template
    )
    return {"status": "queued", "task_id": task.id}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
